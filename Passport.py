import contextlib
import sys
import os
from docx2pdf import convert
from docx import Document
from docx.shared import Pt, RGBColor
from datetime import datetime
from data_update import *
import configparser
import re
import tkinter as tk

# Общая логика работы программы
# создаем интерфейс с кнопками и полями, привязываем к кнопке функцию обработке полей и запуск основной программы
# определяется используемый шаблон паспорта (выбирался на начальном поле),
# далее этот docx шаблон подцепляется программой и в нем будут происходить все изменения (размеченные поля заменяются
# на характеристики продукции)
# Далее введенный артикул с помощью регулярных выражений соотносится к серии продукции и на основе этой серии
# загружается заранее созданный вручную конфигурационный фаил с основными характеристиками серии
# Т.к. внутри серии характеристики могут отличаться разных артикулов, сделаны специальные функции обработчики и
# выделены в отдельный фаил. Результат их работы - обновленный словарь, где ключами служат метки в docx и конфигураторе,
# а значения - то что должно быть вписано вместо меток. Метки одинаковы в обоих фаилах
# далее просто обновляем загруженный в начале docx шаблон и производим все замены внутри фаила по циклу
# согласно обновленному словарю, есть функции форматирования текста

# Известные уязвимости:
# Отсутствуют какие-либо проверки на кривость введенных выражений, подразумевается, что артикулы будут 100% корректны
# и взяты из CRM. Не писал, ибо нет задачи исправлять пользователей, актуально максимально расширить конфиг. фаилы

# создание глобальных переменных для работы с полями интерфейса
name = ""
passport_number = ""
executor = ""
company = ""
data = []
num_lines = 0
count = ""


def keypress(event):
    if event.keycode == 86 and event.keysym != 'v':
        event.widget.event_generate('<<Paste>>')
        return "break"
    elif event.keycode == 67 and event.keysym != 'c':
        event.widget.event_generate('<<Copy>>')
        return "break"
    elif event.keycode == 88 and event.keysym != 'x':
        event.widget.event_generate('<<Cut>>')
        return "break"


#подавление вывода консоли для exe
def suppress_stdout_stderr(func, *args, **kwargs):
    with open(os.devnull, 'w') as devnull:
        with contextlib.redirect_stdout(devnull), contextlib.redirect_stderr(devnull):
            return func(*args, **kwargs)

# Функция для определения директории фаила (exe фаил плохо определяет сам созданные им docx,
# а иначе не получается сконвертировать в pdf
def get_current_directory():
    # Проверка, находимся ли мы в frozen среде (то есть, запущены как EXE)
    if getattr(sys, 'frozen', False):
        # Если да, возвращаем директорию, где находится EXE-файл
        return os.path.dirname(sys.executable)
    else:
        # Иначе возвращаем директорию, где находится исходный Python-файл
        return os.path.dirname(os.path.abspath(__file__))

def replace_text(run, old_text, new_text):
    if old_text in run.text:
        run.text = run.text.replace(old_text, new_text)


# Форматирование текста в правый край (для правой части таблицы)
def format_text_onright(paragraph):
    for run in paragraph.runs:
        if run.text != "":
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 0)
            paragraph.alignment = 2


# Форматирование текста в левый край (для левой части таблицы)
def format_text_onleft(paragraph):
    for run in paragraph.runs:
        if run.text != "":
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0, 0, 0)
            paragraph.alignment = 0


#проверка пустой ячейки
def is_row_empty(row):
    for cell in row.cells:
        if cell.text.strip():
            return False
    return True

#удаление пустых строк в таблице
def delete_rows_with_empty_cells(document):
    tables = document.tables

    for table in tables:
        rows_to_delete = []

        for row in table.rows:
            if is_row_empty(row):
                rows_to_delete.append(row)

        for row in rows_to_delete:
            table._tbl.remove(row._tr)


# функция обработчик нажатия на кнопку
def get_input():
    # объявление глобальных переменных (для считывания с окна ввода)
    global name, passport_number, executor, selected_company, count
    name = entry_name.get()
    passport_number = entry_passport.get()
    executor = entry_executor.get()
    count = entry_count.get()
    # запуск основного кода программы
    run_main_code()


def run_main_code():
    # Получаем значение директории, где исполняется программа (для открытия шаблонов, конфигов и записи в pdf)
    current_dir = get_current_directory()
    # Загружаем существующий шаблон .docx
    if selected_company.get() == "SMC" and checkbox_stamp.get():
        full_path_doc = os.path.join(current_dir, 'templates', 'template_SMCpart_stamp.docx')
    elif selected_company.get() == "SMC" and not checkbox_stamp.get():
        full_path_doc = os.path.join(current_dir, 'templates', 'template_SMCpart.docx')
    elif selected_company.get() == "Indutech" and checkbox_stamp.get():
        full_path_doc = os.path.join(current_dir, 'templates', 'template_INDUTECHpart_stamp.docx')
    elif selected_company.get() == "Indutech" and not checkbox_stamp.get():
        full_path_doc = os.path.join(current_dir, 'templates', 'template_INDUTECHpart.docx')

    document = Document(full_path_doc)
    # Список шаблонов и соответствующих им имен конфигов
    patterns_and_names = {
        r'KFG.*XNT01': 'KFG_XNT01.txt',
        r'KFG.*XRT01': 'KFG_XRT01.txt',
        r'KFG.*XNC01': 'KFG_XNC01.txt',
        r'KQG2.*XKV01': 'KQG2_XKV01.txt',
        r'KQG2.*XLN01': 'KQG2_XLN01.txt',
        r'KQG2.*XRT01': 'KQG2_XRT01.txt',
        r'KQB2.*XRT01': 'KQB2_XRT01.txt',
        r'KQB2.*XLN01': 'KQB2_XLN01.txt',
        r'KQ2.*XLC01': 'KQ2_U_XLC01.txt',
        r'KC.\d{2}-.{2}-XRT01': 'KC_XRT01.txt',
        r'KS.\d{2}-.{2,3}-XLC01': 'KS_XLC01.txt',
        r'K.{2}\d{2}-.{2}-XRT01': 'KS_KX_XRT01.txt',
        r'K.{2}\d{2}-XRT01': 'KP_KS_XRT01.txt',
        r'^H\d{2}-\d{2}-XRT01': 'H_D_XRT01.txt',
        r'^D.\d{2}-\d{2}-XRT01': 'H_D_XRT01.txt',
        r'KQ2.*XRT01': 'KQ2_XRT01.txt',
        r'KJS.*XRT02': 'KQ_XRT02.txt',
        r'KQ.*XRT02': 'KQ_XRT02.txt',
        r'K.F.\d{2}.*XRT01': 'K_F_XRT01.txt',
        r'KPR.*XJC01': 'KPR_XJC01.txt',
        r'KQ.*XLN01': 'KQ_XLN01.txt',
        r'IP3\d{2}': 'IP300.txt',
        r'AW20.*BG-2-L-XLC01': 'AW20_L_XLC01.txt',
        r'AW\d{3}S.*XSP01': 'AW_S_XSP01.txt',
        r'AWH\d{4}.*XLN01': 'AWH_XLN01.txt',
        r'AW\d{4}.*XLN01': 'AW_XLN01.txt',
        r'AF\d{4}.*XLN01': 'AF_XLN01.txt',
        r'AR\d{4}.*XLN01': 'AR_XLN01.txt',
        r'AL\d{4}.*XLN01': 'AL_XLN01.txt',
        r'VHS\d{4}.*XLN01': 'VHS_XLN01.txt',
        r'AC\d{4}.*XLN01': 'AC_XLN01.txt',
        r'AW\d{4}.*XRT01': 'AW_XRT01.txt',
        r'AF\d{4}.*XRT01': 'AF_XRT01.txt',
        r'AR\d{4}.*XRT01': 'AR_XRT01.txt',
        r'AL\d{4}.*XRT01': 'AL_XRT01.txt',
        r'VHS\d{4}.*XRT01': 'VHS_XRT01.txt',
        r'AC\d{4}.*XRT01': 'AC_XRT01.txt',
        r'AW\d{4}.*X425-XNT01': 'AW_X425_XNT01.txt',
        r'AF\d{4}.*X425-XNT01': 'AF_X425_XNT01.txt',
        r'AR\d{4}.*X425-XNT01': 'AR_X425_XNT01.txt',
        r'AP100-\d{2}B-XBR01': 'AP100_XBR01.txt',
        r'AR\d{1}25.*XLN01': 'AR_25_XLN01.txt',
        r'ARH\d{4}.*XLN01': 'ARH_XLN01.txt',
        r'ARH\d{4}.*XMP01': 'ARH_XMP01.txt',
        r'IR\d{4}.*XLN01': 'IR_XLN01.txt',
        r'IR\d{4}.*XRT01': 'IR_XRT01.txt',
        r'ITV2090-.*XTP01': 'ITV2090_XTP01.txt',
        r'ITV\d{4}-.*XKV01': 'ITV_XKV01.txt',
        r'ITV\d{4}-.*XTP01': 'ITV_XTP01.txt',
        r'ITVX2030-.*XTP01': 'ITVX_XTP01.txt',
        r'AFH\d{4}.*XLN01': 'AFH_XLN01.txt',
        r'AFH\d{4}.*XNT01': 'AFH_XNT01.txt',
        r'AFM\d{2}.*XKV01': 'AFM_XKV01.txt',
        r'AFD\d{2}.*XKV01': 'AFD_XKV01.txt',
        r'AMG\d{3}.*XKV01': 'AMG_XKV01.txt',
        r'AFF.*XKV01': 'AFF_XKV01.txt',
        r'AM\d{3}.*XKV01': 'AM_XKV01.txt',
        r'AMD\d{3}.*XKV01': 'AMD_XKV01.txt',
        r'AMH\d{3}.*XKV01': 'AMH_XKV01.txt',
        r'AME\d{3}.*XKV01': 'AME_XKV01.txt',
        r'AMF\d{3}.*XKV01': 'AMF_XKV01.txt',
        r'SY\d{2}20.*XBR01': 'SY_20_XBR01.txt',
        r'SY\d{2}20.*XRT01': 'SY_20_XRT01.txt',
        r'SY\d{2}20.*RU01': 'SY_20_RU01.txt',
        r'^T\d{4}[A-Za-z]{1,2}-.*XRT01': 'T_XRT01.txt',
        r'^TS\d{4}[A-Za-z]{1,2}-.*XLN01': 'TS_XLN01.txt',
        r'^TS\d{4}[A-Za-z]{1,2}-.*XNV0\d{1}': 'TS_XNV01(2).txt',
        r'^TU\d{4}[A-Za-z]{1,2}-.*XRT01': 'TU_XRT01.txt',
        r'^TU\d{4}[A-Za-z]{1,2}-.*XLN01': 'TU_XLN01.txt',
        r'^TI?UB?\d{2,4}[A-Za-z]{1,2}-.*XNV01': 'TU_XNV01.txt',
        r'^TCU\d{4}[A-Za-z]{1,2}-.*XNV01': 'TCU_XNV01.txt',
        r'^STU\d{4}[A-Za-z]{1,2}-.*XLN01': 'STU_XLN01.txt',
        r'^TPFA\d{4}[A-Za-z]{1,2}-.*XKV01': 'TPFA_XKV01.txt',
        r'^TI?PTFE\d{2,4}-.*XNV01': 'TPTFE_XNV01.txt',
        r'^TRBU\d{4}[A-Za-z]{1,2}-.*XKV01': 'TRBU_XKV01.txt',
        r'^TRBU\d{4}[A-Za-z]{1,2}-.*XNV01': 'TRBU_XNV01.txt',
        r'^TRTU\d{4}[A-Za-z]{1,2}-.*XNV01': 'TRTU_XNV01.txt',
        r'^TAU\d{4}[A-Za-z]{1,2}-.*XNV01': 'TAU_XNV01.txt',
        r'^TUDL\d{4}[A-Za-z]{1,2}-.*XNV01': 'TUDL_XNV01.txt',

    }

    # Переменная для хранения имени конфигурации
    config_name = ''

    # Проверка соответствия и установка имени конфигурации
    for pattern, config in patterns_and_names.items():
        if re.match(pattern, name):
            config_name = config
            break
    else:
        output_text.delete('1.0', tk.END)  # Очищаем поле вывода перед новой записью
        return output_text.insert(tk.END, f"Артикул '{name}' не соответствует ни одному из шаблонов.")

    config = configparser.ConfigParser()
    # Указываем разделитель между ключами и значениями
    config.optionxform = str
    # Читаем файл, конфигурации лежат в отдельной папке configs
    #current_dir = get_current_directory()
    full_path_config = os.path.join(current_dir, 'configs', config_name)
    try:
        with open(full_path_config, encoding='utf-8') as file:
            config.read_file(file)
    except FileNotFoundError:
        output_text.delete('1.0', tk.END)  # Очищаем поле вывода перед новой записью
        return output_text.insert(tk.END, f"Файл конфигурации {config_name} не найден.")

    # Получаем значения из секции DEFAULT
    try:
        data = dict(config['DEFAULT'])
    except KeyError:
        output_text.delete('1.0', tk.END)  # Очищаем поле вывода перед новой записью
        return output_text.insert(tk.END, "Секция DEFAULT не найдена в файле конфигурации.")

    num_lines = int(data['Number_lines'])

    if config_name == 'KQ_XLN01.txt':
        # Здесь странный код, суть - нужно обновить словарь и кол-во строк для вывода
        # словарь data обновляется нормально (Mutable объект), а счетчик строк (num_lines) обновляться не хочет
        # теоретически, можно поиграть с объявлением global, но возникают конфликты. Поэтому сделал так:
        # сама функция при работе обновляет словарь и возвращает в довесок обновленное кол-во строк, если меняется.
        # Не везде будет меняться, но для унификации логика будет везде одна и та же
        num_lines = data_update_kqxln01(name, data, num_lines)
    elif config_name == 'AW_XLN01.txt':
        num_lines = data_update_awxln01(name, data, num_lines)
    elif config_name == 'AF_XLN01.txt':
        num_lines = data_update_afxln01(name, data, num_lines)
    elif config_name == 'AR_XLN01.txt':
        num_lines = data_update_arxln01(name, data, num_lines)
    elif config_name == 'AL_XLN01.txt':
        num_lines = data_update_alxln01(name, data, num_lines)
    elif config_name == 'VHS_XLN01.txt':
        num_lines = data_update_vhsxln01(name, data, num_lines)
    elif config_name == 'AC_XLN01.txt':
        num_lines = data_update_acxln01(name, data, num_lines)
    elif config_name == 'AW20_L_XLC01.txt':
        num_lines = data_update_aw20lxlc01(name, data, num_lines)
    elif config_name == 'AW_S_XSP01.txt':
        num_lines = data_update_awsxsp01(name, data, num_lines)
    elif config_name == 'AWH_XLN01.txt':
        num_lines = data_update_awhxln01(name, data, num_lines)
    elif config_name == 'IP300.txt':
        num_lines = data_update_ip300(name, data, num_lines)
    elif config_name == 'AW_XRT01.txt':
        num_lines = data_update_awxrt01(name, data, num_lines)
    elif config_name == 'AF_XRT01.txt':
        num_lines = data_update_afxrt01(name, data, num_lines)
    elif config_name == 'AR_XRT01.txt':
        num_lines = data_update_arxrt01(name, data, num_lines)
    elif config_name == 'AL_XRT01.txt':
        num_lines = data_update_alxrt01(name, data, num_lines)
    elif config_name == 'VHS_XRT01.txt':
        num_lines = data_update_vhsxrt01(name, data, num_lines)
    elif config_name == 'AC_XRT01.txt':
        num_lines = data_update_acxrt01(name, data, num_lines)
    elif config_name == 'AFM_XKV01.txt':
        num_lines = data_update_afmxkv01(name, data, num_lines)
    elif config_name == 'AFD_XKV01.txt':
        num_lines = data_update_afdxkv01(name, data, num_lines)
    elif config_name == 'AW_X425_XNT01.txt':
        num_lines = data_update_awx425xnt01(name, data, num_lines)
    elif config_name == 'AF_X425_XNT01.txt':
        num_lines = data_update_afx425xnt01(name, data, num_lines)
    elif config_name == 'AR_X425_XNT01.txt':
        num_lines = data_update_arx425xnt01(name, data, num_lines)
    elif config_name == 'AP100_XBR01.txt':
        num_lines = data_update_ap100xbr01(name, data, num_lines)
    elif config_name == 'AR_25_XLN01.txt':
        num_lines = data_update_ar_25xln01(name, data, num_lines)
    elif config_name == 'ARH_XLN01.txt':
        num_lines = data_update_arhxln01(name, data, num_lines)
    elif config_name == 'ARH_XMP01.txt':
        num_lines = data_update_arhxmp01(name, data, num_lines)
    elif config_name == 'IR_XLN01.txt':
        num_lines = data_update_irxln01(name, data, num_lines)
    elif config_name == 'IR_XRT01.txt':
        num_lines = data_update_irxrt01(name, data, num_lines)
    elif config_name == 'ITV_XKV01.txt':
        num_lines = data_update_itvxkv01(name, data, num_lines)
    elif config_name == 'ITV_XTP01.txt':
        num_lines = data_update_itvxtp01(name, data, num_lines)
    elif config_name == 'ITVX_XTP01.txt':
        num_lines = data_update_itvxxtp01(name, data, num_lines)
    elif config_name == 'ITV2090_XTP01.txt':
        num_lines = data_update_itv2090xtp01(name, data, num_lines)
    elif config_name == 'AFH_XLN01.txt':
        num_lines = data_update_afhxln01(name, data, num_lines)
    elif config_name == 'AFH_XNT01.txt':
        num_lines = data_update_afhxnt01(name, data, num_lines)
    elif config_name == 'AMG_XKV01.txt':
        num_lines = data_update_amgxkv01(name, data, num_lines)
    elif config_name == 'AFF_XKV01.txt':
        num_lines = data_update_affxkv01(name, data, num_lines)
    elif config_name == 'AM_XKV01.txt':
        num_lines = data_update_amxkv01(name, data, num_lines)
    elif config_name == 'AMD_XKV01.txt':
        num_lines = data_update_amdxkv01(name, data, num_lines)
    elif config_name == 'AMH_XKV01.txt':
        num_lines = data_update_amhxkv01(name, data, num_lines)
    elif config_name == 'AME_XKV01.txt':
        num_lines = data_update_amexkv01(name, data, num_lines)
    elif config_name == 'AMF_XKV01.txt':
        num_lines = data_update_amfxkv01(name, data, num_lines)
    elif re.match(r'SY\d{2}20.*', name):
        num_lines = data_update_sy_20(name, data, num_lines)
    elif config_name == 'KQG2_XRT01.txt':
        num_lines = data_update_kqg2xrt01(name, data, num_lines)
    elif config_name == 'KQG2_XLN01.txt':
        num_lines = data_update_kqg2xln01(name, data, num_lines)
    elif config_name == 'KQG2_XKV01.txt':
        num_lines = data_update_kqg2xkv01(name, data, num_lines)
    elif config_name == 'KFG_XNC01.txt':
        num_lines = data_update_kfgxnc01(name, data, num_lines)
    elif config_name == 'KFG_XRT01.txt':
        num_lines = data_update_kfgxrt01(name, data, num_lines)
    elif config_name == 'KFG_XNT01.txt':
        num_lines = data_update_kfgxnt01(name, data, num_lines)
    elif config_name == 'KQ_XRT02.txt':
        num_lines = data_update_kqxrt02(name, data, num_lines)
    elif config_name == 'KQ2_XRT01.txt':
        num_lines = data_update_kq2xrt01(name, data, num_lines)
    elif config_name == 'KQ2_U_XLC01.txt':
        num_lines = data_update_kq2uxlc01(name, data, num_lines)
    elif config_name == 'KQB2_XRT01.txt':
        num_lines = data_update_kqb2xrt01(name, data, num_lines)
    elif config_name == 'KQB2_XLN01.txt':
        num_lines = data_update_kqb2xln01(name, data, num_lines)
    elif config_name == 'KP_KS_XRT01.txt':
        num_lines = data_update_kpksxrt01(name, data, num_lines)
    elif config_name == 'KS_KX_XRT01.txt':
        num_lines = data_update_kskxxrt01(name, data, num_lines)
    elif config_name == 'KS_XLC01.txt':
        num_lines = data_update_ksxlc01(name, data, num_lines)
    elif config_name == 'KC_XRT01.txt':
        num_lines = data_update_kcxrt01(name, data, num_lines)
    elif config_name == 'KPR_XJC01.txt':
        num_lines = data_update_kprxjc01(name, data, num_lines)
    elif config_name == 'K_F_XRT01.txt':
        num_lines = data_update_kfxrt01(name, data, num_lines)
    elif config_name == 'H_D_XRT01.txt':
        num_lines = data_update_hdxrt01(name, data, num_lines)
    elif config_name == 'T_XRT01.txt':
        num_lines = data_update_txrt01(name, data, num_lines)
    elif config_name == 'TS_XLN01.txt':
        num_lines = data_update_tsxln01(name, data, num_lines)
    elif config_name == 'TS_XNV01(2).txt':
        num_lines = data_update_tsxnv012(name, data, num_lines)
    elif config_name == 'TU_XRT01.txt':
        num_lines = data_update_tuxrt01(name, data, num_lines)
    elif config_name == 'TU_XLN01.txt':
        num_lines = data_update_tuxln01(name, data, num_lines)
    elif config_name == 'TU_XNV01.txt':
        num_lines = data_update_tuxnv01(name, data, num_lines)
    elif config_name == 'TPE_XNV01.txt':
        num_lines = data_update_tpexnv01(name, data, num_lines)
    elif config_name == 'TCU_XNV01.txt':
        num_lines = data_update_tcuxnv01(name, data, num_lines)
    elif config_name == 'STU_XLN01.txt':
        num_lines = data_update_stuxln01(name, data, num_lines)
    elif config_name == 'TPFA_XKV01.txt':
        num_lines = data_update_tpfaxkv01(name, data, num_lines)
    elif config_name == 'TPTFE_XNV01.txt':
        num_lines = data_update_tptfexnv01(name, data, num_lines)
    elif config_name == 'TRBU_XKV01.txt':
        num_lines = data_update_trbuxkv01(name, data, num_lines)
    elif config_name == 'TRBU_XNV01.txt':
        num_lines = data_update_trbuxnv01(name, data, num_lines)
    elif config_name == 'TRTU_XNV01.txt':
        num_lines = data_update_trtuxnv01(name, data, num_lines)
    elif config_name == 'TAU_XNV01.txt':
        num_lines = data_update_tauxnv01(name, data, num_lines)
    elif config_name == 'TUDL_XNV01.txt':
        num_lines = data_update_tudlxnv01(name, data, num_lines)



    # Ниже пошла обработка шаблона и замена полей
    for paragraph in document.paragraphs:
        if 'Declaration' in paragraph.text:
            for run in paragraph.runs:
                if data['Declaration'] == "-":
                    replace_text(run, run.text, "")
                else:
                    dec_string = (f"Продукция имеет\nДекларацию о соответствии\n{data['Declaration']}\n"
                                  f"сроком с {data['DS']} по {data['DE']}")
                    replace_text(run, run.text, dec_string)
                    run.font.name = 'Arial Narrow'
                    run.font.size = Pt(12)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    paragraph.alignment = 0
                    run.font.bold = True
                    break

    for paragraph in document.paragraphs:
        if 'Maker' in paragraph.text:
            for run in paragraph.runs:
                # Получаем текущее время
                current_time = datetime.now()
                # Извлекаем год, месяц и день
                year = current_time.year
                month = f"{current_time.month:02}" # Форматирование месяца до 2 знаков
                day = f"{current_time.day:02}" # Форматирование дня до 2 знаков
                new_string = f"№{year}.{month}.{day}\\{executor}\\{passport_number} от {day}.{month}.{year}"
                replace_text(run, run.text, new_string)
                break

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if 'Name_product' in run.text:
                            replace_text(run, run.text, data['Name_product'])
                            format_text_onright(paragraph)
                            break
                        if 'name' in run.text:
                            replace_text(run, run.text, name)
                            format_text_onright(paragraph)
                            break
                        if 'count' in run.text:
                            replace_text(run, run.text, count)
                            format_text_onright(paragraph)
                            break

                        for i in range(1, num_lines + 1):
                            line_key = f'line{i}'
                            value_key = f'value{i}'
                            if line_key in run.text:
                                replace_text(run, run.text, data[line_key])
                                format_text_onleft(paragraph)
                                break
                            if value_key in run.text:
                                replace_text(run, run.text, data[value_key])
                                format_text_onright(paragraph)
                                break

                        # удаление неиспользуемых ячеек(очистка от значений). В идеале в таблице просто удалять из нее
                        # все лишние строки, но придумать как это сделать не удалось. Думал сформировать таблицу
                        # силами программы и в нее вставлять, но как то все очень криво работало, сделал такой вариант
                        for i in range(num_lines + 1, 9):
                            line_key = f'line{i}'
                            value_key = f'value{i}'
                            if line_key in run.text:
                                replace_text(run, run.text, '')
                                break
                            if value_key in run.text:
                                replace_text(run, run.text, '')
                                break

    #удаление пустых строк в таблице
    delete_rows_with_empty_cells(document)
    # Получаем правильную директорию запущенного фаила (иначе exe фаил может некорректно работать)
    #current_dir = get_current_directory()

    if selected_format.get() == "docx":
        # Сохраняем документ в текущей директории
        full_path_docx = os.path.join(current_dir, name + '.docx')
        document.save(full_path_docx)
    elif selected_format.get() == "pdf":
        # Сохраняем временный документ в текущей директории
        full_path_docx = os.path.join(current_dir, name + '.docx')
        document.save(full_path_docx)
        # Конвертируем docx в pdf
        # full_path_pdf = os.path.join(current_dir, name + '.pdf')
        full_path_pdf = name + '.pdf'
        # Надстройка для перехвата потока вывода записи прогресса в консоль функцией convert, для того,
        # чтобы можно было создать exe фаил с параметром --noconsole, а иначе некрасиво
        suppress_stdout_stderr(convert, full_path_docx, full_path_pdf)
        # Удаляем временный файл
        os.remove(full_path_docx)


    output_text.delete('1.0', tk.END)  # Очищаем поле вывода перед новой записью
    output_text.insert(tk.END, 'Программа успешно отработала!\n')  # Сообщение об успехе

    # проверка на чекбокс, просто if checkbox_out не срабатывает
    if checkbox_out.get():
        for i in range(1, num_lines + 1):
            line_key = f'line{i}'
            value_key = f'value{i}'
            output_text.insert(tk.END, f'{data[line_key]}:   {data[value_key]}\n')


# Создаем главное окно приложения
root = tk.Tk()
#переопределение вставки, т.к. проблема с русской раскладкой
root.bind_all("<Control-KeyPress>", keypress)
root.title("Ввод данных")

# Переменные для хранения выбранного значения радиокнопок
selected_company = tk.StringVar()
selected_company.set("SMC")  # По умолчанию выбираем SMC

selected_format = tk.StringVar()
selected_format.set("docx")  # По умолчанию выбираем docx

# Чекбокс
checkbox_stamp = tk.BooleanVar(value=False)  # Переменная для состояния чекбокса печати
checkbox_out = tk.BooleanVar(value=False)  # Переменная для состояния чекбокса вывода


# Поле для ввода артикула
label_name = tk.Label(root, text="Артикул:")
label_name.grid(row=0, column=0, sticky='w', pady=(10, 0))
entry_name = tk.Entry(root, width=30)
entry_name.grid(row=0, column=1, sticky='w', pady=(10, 0))

# Поле для ввода номера паспорта
label_passport = tk.Label(root, text="Номер паспорта:")
label_passport.grid(row=1, column=0, sticky='w', pady=(3, 0))
entry_passport = tk.Entry(root, width=30)
entry_passport.grid(row=1, column=1, sticky='w', pady=(3, 0))

# Поле для ввода исполнителя
label_executor = tk.Label(root, text="Исполнитель:")
label_executor.grid(row=2, column=0, sticky='w', pady=(3, 0))
entry_executor = tk.Entry(root, width=30)
entry_executor.grid(row=2, column=1, sticky='w', pady=(3, 0))

# Поле для ввода количества
label_count = tk.Label(root, text="Количество:")
label_count.grid(row=3, column=0, sticky='w', pady=(3, 0))
entry_count = tk.Entry(root, width=30)
entry_count.grid(row=3, column=1, sticky='w', pady=(3, 0))

# Радиокнопки для выбора компании
radio_frame = tk.Frame(root)
radio_frame.grid(row=4, columnspan=3, sticky='ew')  # Растягиваем фрейм на три колонки

label_company = tk.Label(radio_frame, text="Шаблон:")
label_company.grid(row=4, column=0, padx=(0, 0), pady=(3, 0))

radio_smc = tk.Radiobutton(radio_frame, text="SMC", variable=selected_company, value="SMC")
radio_smc.grid(row=4, column=1, padx=(30, 20), pady=(3, 0))

radio_indutech = tk.Radiobutton(radio_frame, text="Индутех", variable=selected_company, value="Indutech")
radio_indutech.grid(row=4, column=2, padx=(0, 0), pady=(3, 0))

# Радиокнопки для выбора формата
radio_frame2 = tk.Frame(root)
radio_frame2.grid(row=5, columnspan=3, sticky='ew')  # Растягиваем фрейм на три колонки

label_format = tk.Label(radio_frame2, text="Формат:")
label_format.grid(row=5, column=0, padx=(0, 0), pady=(3, 0))

radio_docx = tk.Radiobutton(radio_frame2, text="docx", variable=selected_format, value="docx")
radio_docx.grid(row=5, column=1, padx=(30, 20), pady=(3, 0))

radio_pdf = tk.Radiobutton(radio_frame2, text="pdf", variable=selected_format, value="pdf")
radio_pdf.grid(row=5, column=2, padx=(0, 0), pady=(3, 0))

# Заголовок для чекбокса
radio_frame3 = tk.Frame(root)
radio_frame3.grid(row=6, columnspan=2, sticky='ew')  # Растягиваем фрейм на две колонки

checkbox_label = tk.Label(radio_frame3, text="Нужна печать исполнителя?")
checkbox_label.grid(row=6, column=0, sticky='w')

checkbox = tk.Checkbutton(radio_frame3, variable=checkbox_stamp)
checkbox.grid(row=6, column=1, sticky='w')

# Заголовок для чекбокса
radio_frame4 = tk.Frame(root)
radio_frame4.grid(row=7, columnspan=2, sticky='ew')  # Растягиваем фрейм на две колонки

checkbox_label = tk.Label(radio_frame4, text="Нужен вывод заполненных тех.характеристик?")
checkbox_label.grid(row=7, column=0, sticky='w')

checkbox = tk.Checkbutton(radio_frame4, variable=checkbox_out)
checkbox.grid(row=7, column=1, sticky='w')

# Поле вывода результата
output_text = tk.Text(root, height=8, width=40)
output_text.grid(row=8, column=0, columnspan=2, pady=(10, 10), padx=(3, 3))

# Кнопка подтверждения
button_confirm = tk.Button(root, text="Подтвердить", command=get_input)
button_confirm.grid(row=9, column=0, columnspan=2, pady=(10, 0))

root.mainloop()

